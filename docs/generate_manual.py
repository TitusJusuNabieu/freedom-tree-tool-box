"""
Generates FreedomTree Field Reporting System — User Manual (.docx)
Run: python3 docs/generate_manual.py
Output: docs/FreedomTree_User_Manual.docx
"""

import io
import os
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Brand colours ────────────────────────────────────────────────────────────
FT_ORANGE   = (238, 91,  77)
FT_GREY_DK  = (55,  70,  85)
FT_GREY_DKR = (16,  20,  25)
FT_GREEN    = (144, 182, 173)
FT_YELLOW   = (246, 185, 79)
FT_GREY_MED = (107, 133, 158)
FT_GREY_LT  = (156, 173, 191)
WHITE       = (255, 255, 255)
PLACEHOLDER_BG = (230, 235, 240)

# ── Helpers ───────────────────────────────────────────────────────────────────

def make_placeholder(width_px: int, height_px: int, label: str, sub: str = "") -> bytes:
    """Return PNG bytes of a labelled screenshot placeholder."""
    img = Image.new("RGB", (width_px, height_px), PLACEHOLDER_BG)
    draw = ImageDraw.Draw(img)

    # Border
    draw.rectangle([0, 0, width_px - 1, height_px - 1],
                   outline=FT_GREY_MED, width=2)

    # Camera icon placeholder (simple circle + rect)
    cx, cy = width_px // 2, height_px // 2 - 30
    r = min(width_px, height_px) // 8
    draw.ellipse([cx - r, cy - r, cx + r, cy + r],
                 outline=FT_GREY_MED, width=3)
    rr = r // 2
    draw.rectangle([cx - rr * 2, cy - r - rr, cx + rr * 2, cy - r + rr // 2],
                   outline=FT_GREY_MED, width=2)

    # Label text (large)
    try:
        font_large = ImageFont.truetype("/Library/Fonts/Arial.ttf", max(14, height_px // 12))
        font_small = ImageFont.truetype("/Library/Fonts/Arial.ttf", max(11, height_px // 18))
    except Exception:
        font_large = ImageFont.load_default()
        font_small = font_large

    bbox = draw.textbbox((0, 0), label, font=font_large)
    tw = bbox[2] - bbox[0]
    draw.text(((width_px - tw) // 2, cy + r + 16), label, fill=FT_GREY_DK, font=font_large)

    if sub:
        bbox2 = draw.textbbox((0, 0), sub, font=font_small)
        tw2 = bbox2[2] - bbox2[0]
        draw.text(((width_px - tw2) // 2, cy + r + 16 + (height_px // 12) + 6),
                  sub, fill=FT_GREY_MED, font=font_small)

    buf = io.BytesIO()
    img.save(buf, format="PNG")
    return buf.getvalue()


def add_placeholder(doc: Document, label: str, sub: str = "",
                    width_in: float = 3.5, height_in: float = 2.4,
                    is_mobile: bool = False):
    """Insert a placeholder image centred in the document."""
    w_px = int(width_in * 96)
    h_px = int(height_in * 96)
    if is_mobile:
        # portrait phone aspect
        w_px = int(1.8 * 96)
        h_px = int(3.2 * 96)
    png = make_placeholder(w_px, h_px, label, sub)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(io.BytesIO(png),
                    width=Inches(1.8 if is_mobile else width_in))
    # Caption
    cap = doc.add_paragraph(f"[ {label} ]")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].font.size = Pt(9)
    cap.runs[0].font.color.rgb = RGBColor(*FT_GREY_MED)
    cap.runs[0].font.italic = True


def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def add_heading(doc: Document, text: str, level: int = 1):
    h = doc.add_heading(text, level=level)
    if level == 1:
        for run in h.runs:
            run.font.color.rgb = RGBColor(*FT_ORANGE)
    elif level == 2:
        for run in h.runs:
            run.font.color.rgb = RGBColor(*FT_GREY_DK)
    return h


def add_note(doc: Document, text: str, colour: tuple = FT_YELLOW):
    """Add a tinted note/tip box."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = "Table Grid"
    cell = tbl.rows[0].cells[0]
    hex_c = "".join(f"{c:02X}" for c in colour)
    set_cell_bg(cell, hex_c)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(*FT_GREY_DKR)
    doc.add_paragraph()   # spacer


def add_step(doc: Document, number: int, title: str, body: str):
    p = doc.add_paragraph()
    run_num = p.add_run(f"Step {number}:  ")
    run_num.bold = True
    run_num.font.color.rgb = RGBColor(*FT_ORANGE)
    run_title = p.add_run(title)
    run_title.bold = True
    if body:
        doc.add_paragraph(body, style="List Continue")


def add_bullet(doc: Document, text: str, level: int = 0):
    p = doc.add_paragraph(text, style="List Bullet" if level == 0 else "List Bullet 2")
    return p


# ── Document assembly ─────────────────────────────────────────────────────────

def build():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # ── Cover page ────────────────────────────────────────────────────────────
    cover_logo = make_placeholder(600, 120, "FREEDOM TREE", "Field Reporting System")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(io.BytesIO(cover_logo), width=Inches(5.0))

    doc.add_paragraph()

    title = doc.add_heading("User Manual", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.color.rgb = RGBColor(*FT_ORANGE)

    sub = doc.add_paragraph("Freedom Tree Field Reporting System\nMobile App & Dashboard")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.size = Pt(14)
    sub.runs[0].font.color.rgb = RGBColor(*FT_GREY_DK)

    doc.add_paragraph()
    ver = doc.add_paragraph("Version 1.0  ·  July 2026")
    ver.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ver.runs[0].font.size = Pt(11)
    ver.runs[0].font.color.rgb = RGBColor(*FT_GREY_MED)

    doc.add_page_break()

    # ── Table of Contents (manual) ────────────────────────────────────────────
    add_heading(doc, "Table of Contents", level=1)
    toc_items = [
        ("1", "Introduction & System Overview"),
        ("2", "Getting Started — Mobile App"),
        ("  2.1", "Installation"),
        ("  2.2", "Logging In"),
        ("  2.3", "Home Screen Overview"),
        ("3", "Recording Data (Mobile)"),
        ("  3.1", "Monthly Health Report"),
        ("  3.2", "Education Survey"),
        ("  3.3", "Maternal Health Survey"),
        ("4", "Syncing Data to the Server"),
        ("5", "Working Offline"),
        ("6", "Profile Management (Mobile)"),
        ("7", "Dashboard — Getting Started"),
        ("  7.1", "Logging In"),
        ("  7.2", "Dashboard Overview"),
        ("8", "Viewing & Managing Reports (Dashboard)"),
        ("9", "User Management (Administrators)"),
        ("10", "Sharing the Dashboard"),
        ("11", "Profile & Account Settings (Dashboard)"),
        ("12", "Roles & Permissions Reference"),
        ("13", "Troubleshooting"),
        ("14", "Contact & Support"),
    ]
    for num, label in toc_items:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(f"{num}   {label}")
        run.font.size = Pt(11)
        if not num.startswith(" "):
            run.bold = True
            run.font.color.rgb = RGBColor(*FT_GREY_DK)
        else:
            run.font.color.rgb = RGBColor(*FT_GREY_MED)

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════════
    # SECTION 1 — INTRODUCTION
    # ═══════════════════════════════════════════════════════════════════════════
    add_heading(doc, "1.  Introduction & System Overview", level=1)
    doc.add_paragraph(
        "The Freedom Tree Field Reporting System is a purpose-built data collection and "
        "monitoring platform for Freedom Tree's community health programmes in Sierra Leone. "
        "It replaces the previous Google Form workflow with a fully integrated system that "
        "works reliably in areas with poor or no internet connectivity."
    )
    doc.add_paragraph()
    add_heading(doc, "What the system includes", level=2)
    add_bullet(doc, "Mobile App (Android) — used by community health field workers to record monthly health reports, education surveys, and maternal health surveys offline, then sync them when connectivity is available.")
    add_bullet(doc, "Web Dashboard — used by Freedom Tree staff and supervisors to view submitted data, run analysis, manage users, and share read-only reports with donors and partners.")
    add_bullet(doc, "Secure server — hosted at Freedom Tree's dedicated server; all data is stored in a PostgreSQL database.")

    doc.add_paragraph()
    add_heading(doc, "Key principles", level=2)
    add_bullet(doc, "Offline-first: data is saved on the device immediately and synced when internet is available. You never lose a completed form.")
    add_bullet(doc, "No duplicates: each record has a unique ID generated on the device. Re-syncing the same record is safe and will never create a duplicate on the server.")
    add_bullet(doc, "Role-based access: field workers only see their own data; supervisors and administrators have broader access.")

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════════
    # SECTION 2 — GETTING STARTED (MOBILE)
    # ═══════════════════════════════════════════════════════════════════════════
    add_heading(doc, "2.  Getting Started — Mobile App", level=1)

    add_heading(doc, "2.1  Installation", level=2)
    doc.add_paragraph(
        "The Freedom Tree app is distributed as an Android APK file. It is not available on "
        "the Google Play Store; you install it directly from the APK provided by your administrator."
    )
    doc.add_paragraph()
    add_step(doc, 1, "Receive the APK",
             'Your administrator will share the file "FreedomTree.apk" via WhatsApp, email, or a shared link.')
    add_step(doc, 2, "Allow installation from unknown sources",
             'On your Android device, go to Settings → Apps → Special app access → Install unknown apps. '
             'Tap your file manager or browser and enable "Allow from this source".')
    add_step(doc, 3, "Open the APK",
             "Tap the downloaded APK file in your notifications or in your Downloads folder. "
             "Tap Install and wait for the installation to complete.")
    add_step(doc, 4, "Open the app",
             "Find the Freedom Tree icon on your home screen and tap it to open the app.")
    add_note(doc,
             "⚠  You only need internet for the very first login. After that, the app works fully offline.",
             FT_YELLOW)

    add_placeholder(doc, "App Icon & Splash Screen", "Orange Freedom Tree logo on launch", is_mobile=True)

    doc.add_paragraph()
    add_heading(doc, "2.2  Logging In", level=2)
    doc.add_paragraph(
        "Your administrator will provide you with a username and password. You must have an "
        "internet connection for your first-ever login."
    )
    add_step(doc, 1, "Enter your username", "Type the username given to you by your administrator.")
    add_step(doc, 2, "Enter your password", "Type your password. Tap the eye icon to show/hide it.")
    add_step(doc, 3, "Tap 'Sign in'",
             "The app will connect to the server, verify your credentials, and take you to the Home screen.")
    add_note(doc,
             "💡  After your first successful login, the app stores your profile securely on the device. "
             "On subsequent starts — even offline — you will be taken straight to the Home screen without needing to log in again.",
             FT_GREEN)

    add_placeholder(doc, "Login Screen", "Username and password fields with Sign In button", is_mobile=True)

    doc.add_paragraph()
    add_heading(doc, "2.3  Home Screen Overview", level=2)
    doc.add_paragraph("The Home screen is divided into four areas:")

    tbl = doc.add_table(rows=5, cols=2)
    tbl.style = "Table Grid"
    headers = ["Area", "Purpose"]
    for i, h in enumerate(headers):
        cell = tbl.rows[0].cells[i]
        set_cell_bg(cell, "EE5B4D")
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(*WHITE)
    rows = [
        ("Welcome card (top)", "Shows your name, community, and position."),
        ("Offline banner", "Appears when you have records waiting to be synced."),
        ("Record data (action cards)", "Tap to start a new Monthly Health Report, Education Survey, or Maternal Health Survey."),
        ("Recent reports", "Shows your previously submitted reports and their sync status (Synced, Pending, Failed)."),
    ]
    for i, (col1, col2) in enumerate(rows):
        tbl.rows[i + 1].cells[0].paragraphs[0].add_run(col1).bold = True
        tbl.rows[i + 1].cells[1].text = col2
    doc.add_paragraph()

    add_placeholder(doc, "Home Screen", "Welcome card, action cards, report list", is_mobile=True)

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════════
    # SECTION 3 — RECORDING DATA
    # ═══════════════════════════════════════════════════════════════════════════
    add_heading(doc, "3.  Recording Data (Mobile)", level=1)
    doc.add_paragraph(
        "All three data entry forms work the same way: fill in the sections, tap Save, and the record "
        "is immediately stored on your device. You do not need internet to save a record."
    )

    add_heading(doc, "3.1  Monthly Health Report", level=2)
    doc.add_paragraph(
        "The Monthly Health Report captures maternal and infant mortality data for one community "
        "in one calendar month. One report per community per month."
    )
    add_step(doc, 1, "Tap 'Monthly health report'", "From the Home screen, tap the orange health report card.")
    add_step(doc, 2, "Select reporting month", "Tap the month selector and choose the month you are reporting for.")
    add_step(doc, 3, "Confirm your community", "Your community is pre-filled from your profile. Edit if needed.")
    add_step(doc, 4, "Fill in the five sections",
             "Work through each section: General, Maternal Health, Infant Health, Contributing Factors, and Narrative. "
             "Required fields are marked with *.")
    add_step(doc, 5, "Tap 'Save report'",
             "The report is saved to your device as Pending and will sync automatically when you are online.")

    add_note(doc,
             "💡  If you already have a report for the same community and month saved on your device, "
             "the app will warn you before creating a duplicate.",
             FT_YELLOW)

    add_placeholder(doc, "Health Report Form — Section 1 (General)",
                    "Reporting month selector, community, submitter details", is_mobile=True)
    add_placeholder(doc, "Health Report Form — Section 2 (Maternal Health)",
                    "Pregnant women count, deliveries, maternal deaths", is_mobile=True)

    doc.add_paragraph()
    add_heading(doc, "Report sections at a glance", level=2)
    tbl2 = doc.add_table(rows=6, cols=2)
    tbl2.style = "Table Grid"
    for i, h in enumerate(["Section", "Fields collected"]):
        cell = tbl2.rows[0].cells[i]
        set_cell_bg(cell, "374655")
        r = cell.paragraphs[0].add_run(h)
        r.bold = True
        r.font.color.rgb = RGBColor(*WHITE)
    section_rows = [
        ("1 — General", "Reporting month, community/district, submitter name & position, date submitted"),
        ("2 — Maternal Health", "Pregnant women, total deliveries (facility/home), maternal deaths, place of death, suspected cause"),
        ("3 — Infant Health", "Live births, infant deaths (within 24 h / 1 month / 12 months)"),
        ("4 — Contributing Factors", "Suspected causes of infant death (multi-select from 5 categories)"),
        ("5 — Narrative", "Key challenges, actions taken/planned, additional comments"),
    ]
    for i, (s, f) in enumerate(section_rows):
        tbl2.rows[i + 1].cells[0].paragraphs[0].add_run(s).bold = True
        tbl2.rows[i + 1].cells[1].text = f
    doc.add_paragraph()

    add_heading(doc, "3.2  Education Survey", level=2)
    doc.add_paragraph(
        "The Education Survey captures a structured assessment at a school or community education setting."
    )
    add_step(doc, 1, "Tap 'Education survey'", "From the Home screen, tap the green education card.")
    add_step(doc, 2, "Enter enumerator details", "Confirm the enumerator name and survey date.")
    add_step(doc, 3, "Enter school/community details", "Type the community or school name and district.")
    add_step(doc, 4, "Enter respondent details", "Provide the respondent's name, category, and sex.")
    add_step(doc, 5, "Answer the survey questions",
             "Work through the questionnaire. Single-choice questions use radio buttons; "
             "rating scale questions (Likert) use a 1–5 row of options.")
    add_step(doc, 6, "Tap 'Save survey'", "The survey is saved locally and queued for sync.")

    add_placeholder(doc, "Education Survey Form", "Survey questions with radio buttons and Likert scale", is_mobile=True)

    doc.add_paragraph()
    add_heading(doc, "3.3  Maternal Health Survey", level=2)
    doc.add_paragraph(
        "The Maternal Health Survey captures detailed pregnancy and birth outcome data at the respondent level."
    )
    add_step(doc, 1, "Tap 'Maternal health survey'", "From the Home screen, tap the yellow maternal card.")
    add_step(doc, 2, "Complete the survey", "Fill in all required fields. Optional fields are labelled as such.")
    add_step(doc, 3, "Tap 'Save survey'", "The survey is saved locally and queued for sync.")

    add_placeholder(doc, "Maternal Health Survey Form", "Community, respondent details, and survey questions", is_mobile=True)

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════════
    # SECTION 4 — SYNCING
    # ═══════════════════════════════════════════════════════════════════════════
    add_heading(doc, "4.  Syncing Data to the Server", level=1)
    doc.add_paragraph(
        "Syncing uploads your locally saved records to the Freedom Tree server so they appear "
        "in the web dashboard. The app syncs automatically — but you can also trigger a manual sync."
    )

    add_heading(doc, "Automatic sync", level=2)
    add_bullet(doc, "When you connect to Wi-Fi or mobile data, the app detects the connection and syncs pending records automatically.")
    add_bullet(doc, "When you return to the app after it has been in the background, it will attempt to sync.")
    add_note(doc, "💡  You do not have to do anything. If you see the Pending badge disappear and the sync icon spin briefly, sync completed successfully.", FT_GREEN)

    add_heading(doc, "Manual sync", level=2)
    doc.add_paragraph(
        "If you want to control exactly which records are uploaded — for example, to upload one specific "
        "report before you lose signal — use the manual sync:"
    )
    add_step(doc, 1, "Tap the sync icon (↻) in the top-right corner of the Home screen",
             "A bottom sheet will appear listing all records waiting to be synced, grouped by type.")
    add_step(doc, 2, "Select the records to upload",
             "All records are pre-selected. Uncheck any you want to hold back. "
             "Tap 'Select all' or 'Deselect all' to quickly toggle everything.")
    add_step(doc, 3, "Tap 'Sync X records'",
             "The sheet switches to a progress view showing an animated progress bar and a live count of uploaded records.")
    add_step(doc, 4, "Wait for completion",
             "When all selected records have uploaded, a green ✓ confirmation appears. "
             "The sheet closes automatically after a moment.")
    add_note(doc, "💡  After syncing, the Recent reports list will update automatically to reflect the new sync status of each record.", FT_GREEN)

    add_placeholder(doc, "Sync Sheet — Selection view",
                    "Grouped checklist of pending reports, Sync button", is_mobile=True)
    add_placeholder(doc, "Sync Sheet — Progress view",
                    "Animated orange progress bar, upload counter", is_mobile=True)
    add_placeholder(doc, "Sync Sheet — Done view",
                    "Green cloud-done icon, record count confirmation", is_mobile=True)

    doc.add_paragraph()
    add_heading(doc, "Sync status badges", level=2)
    tbl3 = doc.add_table(rows=6, cols=2)
    tbl3.style = "Table Grid"
    for i, h in enumerate(["Badge", "Meaning"]):
        cell = tbl3.rows[0].cells[i]
        set_cell_bg(cell, "374655")
        r = cell.paragraphs[0].add_run(h)
        r.bold = True
        r.font.color.rgb = RGBColor(*WHITE)
    badge_rows = [
        ("Pending", "Saved on device, waiting to be uploaded."),
        ("Syncing…", "Currently being uploaded to the server."),
        ("Synced", "Successfully received by the server."),
        ("Failed", "Upload failed (e.g. server error). The record stays safely on your device and will retry next sync."),
        ("Draft", "Partially filled form that has not been submitted yet."),
    ]
    for i, (b, m) in enumerate(badge_rows):
        tbl3.rows[i + 1].cells[0].paragraphs[0].add_run(b).bold = True
        tbl3.rows[i + 1].cells[1].text = m
    doc.add_paragraph()

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════════
    # SECTION 5 — OFFLINE USAGE
    # ═══════════════════════════════════════════════════════════════════════════
    add_heading(doc, "5.  Working Offline", level=1)
    doc.add_paragraph(
        "The Freedom Tree app is designed for environments where internet connectivity is unreliable. "
        "Here is what works with and without internet:"
    )
    tbl4 = doc.add_table(rows=7, cols=3)
    tbl4.style = "Table Grid"
    for i, h in enumerate(["Feature", "Offline", "Online"]):
        cell = tbl4.rows[0].cells[i]
        set_cell_bg(cell, "EE5B4D")
        r = cell.paragraphs[0].add_run(h)
        r.bold = True
        r.font.color.rgb = RGBColor(*WHITE)
    offline_rows = [
        ("Log in (first time)", "✗  Requires internet", "✓"),
        ("Log in (after first time)", "✓  Uses cached profile", "✓"),
        ("Fill in a report/survey", "✓", "✓"),
        ("Save a report/survey", "✓  Saved locally", "✓"),
        ("Sync records to server", "✗  Queued until online", "✓  Automatic or manual"),
        ("View recent records on device", "✓", "✓"),
    ]
    for i, (f, off, on) in enumerate(offline_rows):
        tbl4.rows[i + 1].cells[0].paragraphs[0].add_run(f)
        tbl4.rows[i + 1].cells[1].paragraphs[0].add_run(off)
        tbl4.rows[i + 1].cells[2].paragraphs[0].add_run(on)
    doc.add_paragraph()
    add_note(doc,
             "⚠  If the orange banner 'X items waiting to sync' appears at the top of the Home screen, "
             "it means you have records that have not yet reached the server. Connect to internet and they will upload automatically.",
             FT_YELLOW)

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════════
    # SECTION 6 — MOBILE PROFILE
    # ═══════════════════════════════════════════════════════════════════════════
    add_heading(doc, "6.  Profile Management (Mobile)", level=1)
    doc.add_paragraph("Your profile stores the name and position that pre-fill on every report you submit.")
    add_step(doc, 1, "Tap the person icon (👤) in the top-right corner",
             "The Profile screen opens.")
    add_step(doc, 2, "Edit your name or position",
             "Tap the Name or Position field and type your changes.")
    add_step(doc, 3, "Update your profile photo",
             "Tap the orange camera button on your avatar to choose a photo from your gallery.")
    add_step(doc, 4, "Tap 'Save changes'",
             "Your updated profile is saved to both the server and your device. "
             "Changes are visible immediately, even if you go offline afterwards.")
    add_step(doc, 5, "Change your password",
             "Tap 'Change password', enter your current password, then your new password twice, and tap 'Update password'.")

    add_placeholder(doc, "Profile Screen", "Avatar with camera overlay, name/position fields, save button", is_mobile=True)

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════════
    # SECTION 7 — DASHBOARD GETTING STARTED
    # ═══════════════════════════════════════════════════════════════════════════
    add_heading(doc, "7.  Dashboard — Getting Started", level=1)
    doc.add_paragraph(
        "The web dashboard is accessed from any modern browser. No installation is required."
    )
    add_note(doc, f"🌐  Dashboard address: http://142.93.246.217:8080", FT_GREY_LT)

    add_heading(doc, "7.1  Logging In", level=2)
    add_step(doc, 1, "Open your browser", "Open Chrome, Firefox, or Edge and navigate to the dashboard address above.")
    add_step(doc, 2, "Enter your credentials", "Type your username and password and click 'Sign in'.")
    add_step(doc, 3, "You are taken to the Dashboard home page", "The home page shows summary statistics and trend charts.")

    add_placeholder(doc, "Dashboard — Login Page",
                    "Username, password fields and Sign in button", width_in=5.5, height_in=3.0)

    doc.add_paragraph()
    add_heading(doc, "7.2  Dashboard Overview", level=2)
    doc.add_paragraph("The dashboard uses a left navigation bar with the following sections:")
    add_bullet(doc, "Dashboard — Summary stat cards and trend charts for the current period.")
    add_bullet(doc, "Reports — Full table of all submitted reports with filters.")
    add_bullet(doc, "Education Surveys — Table of submitted education questionnaires.")
    add_bullet(doc, "Maternal Surveys — Table of submitted maternal questionnaires.")
    add_bullet(doc, "Settings → Users — (Admins only) Manage user accounts.")
    add_bullet(doc, "Settings → Profile — Update your name, avatar, or password.")

    add_placeholder(doc, "Dashboard — Home Page",
                    "Stat cards: maternal deaths, live births, deliveries; trend chart", width_in=5.5, height_in=3.2)

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════════
    # SECTION 8 — VIEWING REPORTS (DASHBOARD)
    # ═══════════════════════════════════════════════════════════════════════════
    add_heading(doc, "8.  Viewing & Managing Reports (Dashboard)", level=1)

    add_heading(doc, "Report list", level=2)
    doc.add_paragraph(
        "Navigate to Reports in the left menu. The table shows all submitted reports. "
        "Use the filters at the top to narrow down by month, community, or date range."
    )
    add_bullet(doc, "Click any row to open the full report detail.")
    add_bullet(doc, "Administrators can edit any field in the detail view.")
    add_bullet(doc, "Use the Export button to download the filtered data as a CSV file.")

    add_placeholder(doc, "Dashboard — Reports Table",
                    "Table with month, community, maternal deaths columns and filter bar", width_in=5.5, height_in=3.0)

    add_heading(doc, "Report detail", level=2)
    doc.add_paragraph(
        "Click a report row to open the full detail view. Administrators can click 'Edit' "
        "to correct any field. All edits are recorded with a timestamp."
    )

    add_placeholder(doc, "Dashboard — Report Detail",
                    "All 15 data fields displayed, with Edit button for admins", width_in=5.5, height_in=3.0)

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════════
    # SECTION 9 — USER MANAGEMENT
    # ═══════════════════════════════════════════════════════════════════════════
    add_heading(doc, "9.  User Management (Administrators)", level=1)
    doc.add_paragraph(
        "Administrators and Super Administrators can manage user accounts from Settings → Users. "
        "Field Workers and Supervisors do not have access to this section."
    )

    add_heading(doc, "Creating a new user", level=2)
    add_step(doc, 1, "Go to Settings → Users", "Click 'Settings' in the left navigation, then 'Users'.")
    add_step(doc, 2, "Click 'New user'", "The new user form opens.")
    add_step(doc, 3, "Fill in details",
             "Enter the full name, username, community (optional), position, and role. "
             "A temporary password will be generated — share it with the user securely.")
    add_step(doc, 4, "Click 'Create user'", "The account is created immediately and the user can log in on the mobile app.")
    add_note(doc,
             "⚠  Administrators can create Field Worker, Supervisor, and Data Analyst accounts. "
             "Only a Super Administrator can create or modify other Administrator accounts.",
             FT_YELLOW)

    add_placeholder(doc, "Dashboard — Users List",
                    "Table of all users with role badges and active status toggle", width_in=5.5, height_in=3.0)
    add_placeholder(doc, "Dashboard — New User Form",
                    "Name, username, role dropdown, community, position fields", width_in=5.5, height_in=3.0)

    add_heading(doc, "Deactivating a user", level=2)
    doc.add_paragraph(
        "To prevent a user from logging in (e.g. a field worker who has left the programme), "
        "find them in the Users list and toggle their Active status to Off. "
        "Their historical data is preserved."
    )

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════════
    # SECTION 10 — SHARING
    # ═══════════════════════════════════════════════════════════════════════════
    add_heading(doc, "10.  Sharing the Dashboard", level=1)
    doc.add_paragraph(
        "You can generate a time-limited, read-only link to share the dashboard with donors, "
        "partners, or board members. The recipient does not need an account."
    )
    add_step(doc, 1, "Click 'Share dashboard' in the top navigation",
             "The share link dialog opens.")
    add_step(doc, 2, "Enter a label",
             "Type a descriptive name for this link (e.g. 'Q2 2026 Donor Report'). "
             "This is for your reference and appears in the shared link list.")
    add_step(doc, 3, "Choose a community scope",
             "Select 'All communities' or restrict the link to one specific community.")
    add_step(doc, 4, "Choose an expiry period", "Select 30, 60, 90 days, 6 months, or 1 year.")
    add_step(doc, 5, "Click 'Generate link'", "A unique read-only URL is created.")
    add_step(doc, 6, "Copy and share the link",
             "Click the copy button (or click the URL to select it and copy manually). "
             "Send the link to the recipient — they can open it in any browser.")
    add_note(doc, "💡  The link expires automatically on the date shown. You can generate multiple links with different scopes and expiry dates.", FT_GREEN)

    add_placeholder(doc, "Dashboard — Share Link Dialog",
                    "Label, community scope, expiry selector, generated URL with copy button", width_in=4.5, height_in=3.2)

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════════
    # SECTION 11 — DASHBOARD PROFILE
    # ═══════════════════════════════════════════════════════════════════════════
    add_heading(doc, "11.  Profile & Account Settings (Dashboard)", level=1)
    doc.add_paragraph("Access your profile from Settings → Profile in the left menu.")
    add_bullet(doc, "Change your display name or position — this is reflected on any reports you submit from the web.")
    add_bullet(doc, "Upload a profile photo — click the avatar area to select an image.")
    add_bullet(doc, "Change your password — enter your current password, then type and confirm a new one.")
    add_note(doc, "💡  Profile changes on the dashboard are independent of the mobile app profile. Update both if you need consistency.", FT_YELLOW)

    add_placeholder(doc, "Dashboard — Profile Page",
                    "Avatar upload, name/position fields, change password section", width_in=5.5, height_in=3.0)

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════════
    # SECTION 12 — ROLES & PERMISSIONS
    # ═══════════════════════════════════════════════════════════════════════════
    add_heading(doc, "12.  Roles & Permissions Reference", level=1)

    tbl5 = doc.add_table(rows=6, cols=5)
    tbl5.style = "Table Grid"
    col_headers = ["Permission", "Field Worker", "Supervisor / Data Analyst", "Administrator", "Super Admin"]
    for i, h in enumerate(col_headers):
        cell = tbl5.rows[0].cells[i]
        set_cell_bg(cell, "374655")
        r = cell.paragraphs[0].add_run(h)
        r.bold = True
        r.font.color.rgb = RGBColor(*WHITE)
        r.font.size = Pt(9)
    perm_rows = [
        ("Submit reports/surveys (mobile)", "✓ Own only", "—", "✓ All", "✓ All"),
        ("View reports (dashboard)", "✗", "✓ All", "✓ All", "✓ All"),
        ("Edit reports (dashboard)", "✗", "✗", "✓ All", "✓ All"),
        ("Manage users", "✗", "✗", "✓ Non-admin only", "✓ All"),
        ("Generate share links", "✗", "✓", "✓", "✓"),
    ]
    for i, row in enumerate(perm_rows):
        for j, val in enumerate(row):
            cell = tbl5.rows[i + 1].cells[j]
            p = cell.paragraphs[0]
            r = p.add_run(val)
            r.font.size = Pt(9)
            if j == 0:
                r.bold = True
    doc.add_paragraph()

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════════
    # SECTION 13 — TROUBLESHOOTING
    # ═══════════════════════════════════════════════════════════════════════════
    add_heading(doc, "13.  Troubleshooting", level=1)

    problems = [
        (
            "The app shows 'Could not reach the server'",
            "This means the app cannot connect to the Freedom Tree server.",
            [
                "Check that you have an active mobile data or Wi-Fi connection.",
                "Try opening a web page in your browser to confirm internet is working.",
                "If internet is working but the error persists, the server may be temporarily unavailable — try again in a few minutes.",
                "Your data is safe on your device. You can continue filling in forms offline.",
            ]
        ),
        (
            "Sync shows 'Failed' badge on a report",
            "The record reached the server but was rejected (for example, a duplicate report for the same community and month already exists).",
            [
                "Tap the report to view the error message.",
                "Contact your administrator if you believe the failure is an error.",
                "The report is still safely stored on your device.",
            ]
        ),
        (
            "I forgot my password",
            "Password reset is done by your administrator.",
            [
                "Contact your supervisor or administrator.",
                "The administrator can set a new temporary password for you in the Users section of the dashboard.",
            ]
        ),
        (
            "The app is very slow",
            "The app may have accumulated many local records.",
            [
                "Connect to Wi-Fi and allow the app to sync all pending records.",
                "If the problem continues, contact your administrator.",
            ]
        ),
        (
            "I submitted a report with an error",
            "Contact your administrator.",
            [
                "Administrators can edit any report field in the dashboard.",
                "Provide your administrator with the community name and reporting month.",
            ]
        ),
        (
            "The share link doesn't work for the recipient",
            "The link may have expired.",
            [
                "Check the expiry date — generate a new link if it has expired.",
                "Make sure the full URL was copied without being cut off.",
                "The recipient should open the link in Chrome, Firefox, or Edge.",
            ]
        ),
    ]

    for problem, cause, steps in problems:
        p = doc.add_paragraph()
        r = p.add_run(f"Problem: {problem}")
        r.bold = True
        r.font.color.rgb = RGBColor(*FT_ORANGE)
        doc.add_paragraph(f"Cause: {cause}")
        for s in steps:
            add_bullet(doc, s)
        doc.add_paragraph()

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════════
    # SECTION 14 — CONTACT & SUPPORT
    # ═══════════════════════════════════════════════════════════════════════════
    add_heading(doc, "14.  Contact & Support", level=1)
    doc.add_paragraph(
        "For technical support with the Freedom Tree Field Reporting System, contact your "
        "supervisor or the system administrator. For issues that cannot be resolved internally:"
    )
    add_bullet(doc, "System administrator email: nabieutitusjusu@gmail.com")
    add_bullet(doc, "Provide: your name, username, device model, and a description of the issue.")
    add_note(doc,
             "📋  If reporting a sync failure, also note the community name, reporting month, "
             "and the error message shown on the report badge.",
             FT_GREY_LT)

    # ── Save ─────────────────────────────────────────────────────────────────
    out_path = os.path.join(os.path.dirname(__file__), "FreedomTree_User_Manual.docx")
    doc.save(out_path)
    print(f"✅  Manual saved to: {out_path}")


if __name__ == "__main__":
    build()
